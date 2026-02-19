"""
User document upload pipeline for CountyIQ.
Auto-detects file type (PDF, CSV, DOCX, TXT, MD) and routes to appropriate processor.
# DP-100: Data pipeline - Multi-format ingestion and routing for RAG data preparation.
"""

import re
from typing import Any

from loguru import logger

from data.schemas.document import ContentType, CountyDocument, DocumentCategory
from data.schemas.registry_loader import get_county
from pipelines.ingest.pdf_processor import PDFProcessor
from pipelines.ingest.structured_processor import StructuredProcessor

# Sentinel FIPS when user does not specify a county
USER_UPLOAD_FIPS = "00000"
USER_UPLOAD_COUNTY = "User Upload"
USER_UPLOAD_STATE = "US"

CHUNK_SIZE = 1000
OVERLAP = 200


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = OVERLAP) -> list[str]:
    """Split text into overlapping chunks (same logic as PDFProcessor)."""
    if overlap >= chunk_size:
        overlap = max(0, chunk_size // 2)
    text = text.strip()
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def _detect_file_type(file_bytes: bytes, filename: str) -> str:
    """
    Detect file type by magic bytes and extension.
    Returns: 'pdf', 'csv', 'docx', 'txt', 'md', or 'unknown'.
    """
    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "").strip()
    # PDF magic
    if file_bytes[:4] == b"%PDF" or ext == "pdf":
        return "pdf"
    # DOCX is a zip
    if file_bytes[:2] == b"PK" and (ext == "docx" or "word" in filename.lower()):
        return "docx"
    # CSV: heuristic
    if ext == "csv":
        return "csv"
    try:
        first_line = file_bytes.split(b"\n")[0].decode("utf-8", errors="ignore")
        if "," in first_line and ext == "csv":
            return "csv"
    except Exception:
        pass
    if ext in ("txt", "text"):
        return "txt"
    if ext == "md":
        return "md"
    return "unknown"


# DP-100: Data pipeline - Multi-format ingestion and routing
class UploadProcessor:
    """
    Process user-uploaded files into CountyDocuments.
    Auto-detects PDF, CSV, DOCX, TXT, MD and routes to the appropriate extractor.
    """

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        overlap: int = OVERLAP,
    ) -> None:
        self.chunk_size = chunk_size
        self.overlap = overlap
        self._pdf_processor = PDFProcessor(chunk_size=chunk_size, overlap=overlap)
        self._structured_processor = StructuredProcessor()

    def extract_docx(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            file_bytes: Raw DOCX file bytes.

        Returns:
            Extracted text, or empty string on failure.
        """
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            text = "\n".join(parts).strip()
            logger.debug("Extracted {} chars from DOCX", len(text))
            return text
        except Exception as e:
            logger.warning("DOCX extraction failed: {}", e)
            return ""

    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """
        Extract text from TXT or MD file.

        Args:
            file_bytes: Raw file bytes.
            filename: Original filename (for encoding hint).

        Returns:
            Decoded text.
        """
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("latin-1", errors="replace")

    def _fips_and_names(self, fips: str | None) -> tuple[str, str, str]:
        """Return (fips, county_name, state_abbr); use placeholders when fips is None."""
        if fips and str(fips).strip().zfill(5) != "00000":
            fips = str(fips).strip().zfill(5)
            county = get_county(fips)
            if county:
                return fips, county.county_name, county.state_abbr
        return USER_UPLOAD_FIPS, USER_UPLOAD_COUNTY, USER_UPLOAD_STATE

    def process(
        self,
        file_bytes: bytes,
        filename: str,
        fips: str | None = None,
        category: DocumentCategory = DocumentCategory.user_upload,
    ) -> list[CountyDocument]:
        """
        Process an uploaded file into CountyDocuments.
        Auto-detects type and routes to PDF, CSV, DOCX, or text chunking.

        Args:
            file_bytes: Raw file bytes.
            filename: Original filename (used for type detection and metadata).
            fips: Optional 5-digit county FIPS; if None, document is not county-specific.
            category: Document category (default user_upload).

        Returns:
            List of CountyDocument instances.
        """
        fips_val, county_name, state_abbr = self._fips_and_names(fips)
        file_type = _detect_file_type(file_bytes, filename)

        if file_type == "pdf":
            return self._process_pdf(file_bytes, filename, fips_val, county_name, state_abbr, category)
        if file_type == "csv":
            return self._process_csv(file_bytes, filename, fips_val, county_name, state_abbr, category)
        if file_type == "docx":
            return self._process_docx(file_bytes, filename, fips_val, county_name, state_abbr, category)
        if file_type in ("txt", "md"):
            return self._process_text(file_bytes, filename, fips_val, county_name, state_abbr, category)

        logger.warning("Unsupported file type for {}: {}", filename, file_type)
        return []

    def _process_pdf(
        self,
        file_bytes: bytes,
        filename: str,
        fips: str,
        county_name: str,
        state_abbr: str,
        category: DocumentCategory,
    ) -> list[CountyDocument]:
        """Route PDF to PDFProcessor; override county info for user upload."""
        docs = self._pdf_processor.process(
            source_url=filename,
            fips=fips,
            category=category,
            pdf_bytes=file_bytes,
            county_name=county_name,
            state_abbr=state_abbr,
        )
        for d in docs:
            d.metadata["original_filename"] = filename
            d.metadata["file_type"] = "pdf"
            d.metadata["upload_source"] = "user"
        return docs

    def _process_csv(
        self,
        file_bytes: bytes,
        filename: str,
        fips: str,
        county_name: str,
        state_abbr: str,
        category: DocumentCategory,
    ) -> list[CountyDocument]:
        """Route CSV to StructuredProcessor; ensure county info is set."""
        docs = self._structured_processor.process_csv(
            source_url=filename,
            fips=fips,
            category=category,
            csv_bytes=file_bytes,
            county_name_override=county_name if fips == USER_UPLOAD_FIPS else None,
            state_abbr_override=state_abbr if fips == USER_UPLOAD_FIPS else None,
        )
        for d in docs:
            d.metadata["original_filename"] = filename
            d.metadata["file_type"] = "csv"
            d.metadata["upload_source"] = "user"
        return docs

    def _process_docx(
        self,
        file_bytes: bytes,
        filename: str,
        fips: str,
        county_name: str,
        state_abbr: str,
        category: DocumentCategory,
    ) -> list[CountyDocument]:
        """Extract text from DOCX and chunk into CountyDocuments."""
        text = self.extract_docx(file_bytes)
        if not text.strip():
            logger.warning("DOCX produced no text: {}", filename)
            return []
        return self._text_to_documents(
            text, filename, fips, county_name, state_abbr, category, "docx"
        )

    def _process_text(
        self,
        file_bytes: bytes,
        filename: str,
        fips: str,
        county_name: str,
        state_abbr: str,
        category: DocumentCategory,
    ) -> list[CountyDocument]:
        """Decode TXT/MD and chunk into CountyDocuments."""
        text = self.extract_text(file_bytes, filename)
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        return self._text_to_documents(
            text, filename, fips, county_name, state_abbr, category, ext
        )

    def _text_to_documents(
        self,
        text: str,
        filename: str,
        fips: str,
        county_name: str,
        state_abbr: str,
        category: DocumentCategory,
        file_type: str,
    ) -> list[CountyDocument]:
        """Chunk text and create one CountyDocument per chunk."""
        chunks = _chunk_text(text, self.chunk_size, self.overlap)
        if not chunks:
            chunks = [text[: self.chunk_size]] if text.strip() else []

        documents: list[CountyDocument] = []
        for i, chunk_content in enumerate(chunks):
            meta: dict[str, Any] = {
                "original_filename": filename,
                "file_type": file_type,
                "upload_source": "user",
                "chunk_index": i,
                "total_chunks": len(chunks),
            }
            doc = CountyDocument(
                fips=fips,
                county_name=county_name,
                state_abbr=state_abbr,
                category=category,
                source_url=filename,
                content_type=ContentType.text,
                raw_content=chunk_content,
                processed_content=chunk_content.strip(),
                metadata=meta,
            )
            documents.append(doc)

        logger.info("Processed {} into {} chunks", filename, len(documents))
        return documents
